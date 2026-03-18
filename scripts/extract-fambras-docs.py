"""
Extrai conteudo de todos os documentos Fambras (PDF, DOCX, XLSX, DOC)
para arquivos de texto puro organizados por fase.
"""
import os
import sys
import traceback

# Libraries
import fitz  # PyMuPDF for PDFs
from docx import Document  # python-docx for DOCX
import openpyxl  # for XLSX

BASE_DIR = r"C:\HalalSphere\Documentos Certificação"
OUTPUT_DIR = r"C:\Projetos\halalsphere-docs\PLANNING\fambras-extraido"

def extract_pdf(filepath):
    """Extract text from PDF using PyMuPDF."""
    doc = fitz.open(filepath)
    text_parts = []
    for i, page in enumerate(doc):
        text = page.get_text()
        text_parts.append(f"=== PAGINA {i+1}/{len(doc)} ===\n{text}")
    doc.close()
    return "\n\n".join(text_parts)

def extract_docx(filepath):
    """Extract text from DOCX including paragraphs and tables."""
    doc = Document(filepath)
    text_parts = []

    # Extract paragraphs
    text_parts.append("=== PARAGRAFOS ===")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            style = para.style.name if para.style else "Normal"
            text_parts.append(f"[{style}] {para.text}")

    # Extract tables
    for t_idx, table in enumerate(doc.tables):
        text_parts.append(f"\n=== TABELA {t_idx+1} ({len(table.rows)} linhas x {len(table.columns)} colunas) ===")
        for r_idx, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                cells.append(cell_text)
            text_parts.append(f"Linha {r_idx+1}: " + " | ".join(cells))

    return "\n".join(text_parts)

def extract_xlsx(filepath):
    """Extract text from XLSX including all sheets."""
    try:
        wb = openpyxl.load_workbook(filepath, data_only=True)
    except Exception as e:
        return f"ERRO ao abrir XLSX: {e}"

    text_parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text_parts.append(f"=== SHEET: {sheet_name} (dimensoes: {ws.dimensions}) ===")

        row_count = 0
        for row in ws.iter_rows(values_only=False):
            vals = []
            for cell in row:
                if cell.value is not None:
                    val = str(cell.value).strip()
                    if val:
                        vals.append(val)
                    else:
                        vals.append("")
                else:
                    vals.append("")
            if any(v for v in vals):
                text_parts.append(" | ".join(vals))
                row_count += 1

        if row_count == 0:
            text_parts.append("(sheet vazia)")

    wb.close()
    return "\n".join(text_parts)

def extract_doc(filepath):
    """For .DOC files - try to extract with basic method."""
    # .DOC (old format) - PyMuPDF cannot read, python-docx cannot read
    # Try antiword or just flag as unreadable
    return f"[FORMATO .DOC LEGADO - nao extraivel automaticamente]\nArquivo: {os.path.basename(filepath)}\nTamanho: {os.path.getsize(filepath)} bytes"

def sanitize_filename(name):
    """Create safe filename for output."""
    # Remove problematic chars
    for ch in ['<', '>', ':', '"', '/', '\\', '|', '?', '*']:
        name = name.replace(ch, '_')
    return name[:200]  # limit length

def get_phase_folder(rel_path):
    """Extract phase folder from relative path."""
    parts = rel_path.split(os.sep)
    if len(parts) >= 1:
        return parts[0]
    return "outros"

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    stats = {"pdf": 0, "docx": 0, "xlsx": 0, "doc": 0, "errors": 0, "total": 0}
    errors = []

    # Walk all files
    all_files = []
    for root, dirs, files in os.walk(BASE_DIR):
        for fname in sorted(files):
            filepath = os.path.join(root, fname)
            rel_path = os.path.relpath(filepath, BASE_DIR)
            all_files.append((filepath, rel_path, fname))

    print(f"Total de arquivos encontrados: {len(all_files)}")

    # Create a single consolidated output per phase
    phase_contents = {}

    for filepath, rel_path, fname in all_files:
        stats["total"] += 1
        ext = fname.rsplit('.', 1)[-1].lower() if '.' in fname else ''
        phase = get_phase_folder(rel_path)

        if phase not in phase_contents:
            phase_contents[phase] = []

        print(f"[{stats['total']}/{len(all_files)}] {rel_path[:80]}...", end=" ")

        try:
            if ext == 'pdf':
                content = extract_pdf(filepath)
                stats["pdf"] += 1
            elif ext == 'docx':
                content = extract_docx(filepath)
                stats["docx"] += 1
            elif ext == 'xlsx':
                content = extract_xlsx(filepath)
                stats["xlsx"] += 1
            elif ext == 'doc':
                content = extract_doc(filepath)
                stats["doc"] += 1
            else:
                content = f"[FORMATO NAO SUPORTADO: .{ext}]"

            header = f"\n{'='*80}\nARQUIVO: {fname}\nCAMINHO: {rel_path}\nTAMANHO: {os.path.getsize(filepath)} bytes\n{'='*80}\n"
            phase_contents[phase].append(header + content)
            print("OK")

        except Exception as e:
            stats["errors"] += 1
            error_msg = f"ERRO: {fname} -> {str(e)}"
            errors.append(error_msg)
            phase_contents[phase].append(f"\n{'='*80}\nARQUIVO: {fname} [ERRO: {e}]\n{'='*80}\n")
            print(f"ERRO: {e}")

    # Write consolidated files per phase
    for phase, contents in sorted(phase_contents.items()):
        safe_name = sanitize_filename(phase)
        output_file = os.path.join(OUTPUT_DIR, f"{safe_name}.txt")
        with open(output_file, 'w', encoding='utf-8', errors='replace') as f:
            f.write(f"FASE: {phase}\n")
            f.write(f"Total de arquivos: {len(contents)}\n")
            f.write(f"{'='*80}\n\n")
            f.write("\n\n".join(contents))
        print(f"\nGravado: {output_file} ({len(contents)} docs)")

    # Write summary
    summary_file = os.path.join(OUTPUT_DIR, "_RESUMO_EXTRACAO.txt")
    with open(summary_file, 'w', encoding='utf-8') as f:
        f.write("RESUMO DA EXTRACAO DE DOCUMENTOS FAMBRAS\n")
        f.write(f"{'='*50}\n")
        f.write(f"Total de arquivos: {stats['total']}\n")
        f.write(f"PDFs extraidos: {stats['pdf']}\n")
        f.write(f"DOCX extraidos: {stats['docx']}\n")
        f.write(f"XLSX extraidos: {stats['xlsx']}\n")
        f.write(f"DOC (legado): {stats['doc']}\n")
        f.write(f"Erros: {stats['errors']}\n")
        f.write(f"\nFases processadas: {len(phase_contents)}\n")
        for phase in sorted(phase_contents.keys()):
            f.write(f"  - {phase}: {len(phase_contents[phase])} docs\n")
        if errors:
            f.write(f"\nERROS:\n")
            for e in errors:
                f.write(f"  {e}\n")

    print(f"\n{'='*50}")
    print(f"CONCLUIDO!")
    print(f"PDFs: {stats['pdf']}, DOCX: {stats['docx']}, XLSX: {stats['xlsx']}, DOC: {stats['doc']}")
    print(f"Erros: {stats['errors']}")
    print(f"Saida em: {OUTPUT_DIR}")

if __name__ == "__main__":
    main()
